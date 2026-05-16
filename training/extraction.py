import logging
from dataclasses import dataclass

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

logger = logging.getLogger(__name__)


class _PositionCursor:
    """Tracks document position and accumulates full_text, enforcing the offset invariant.

    Every call to consume(text) advances the cursor by len(text)+1 (the +1 is the
    newline separator that will appear in build_full_text()), so full_text[start:end]
    == text is guaranteed by construction for every (start, end) pair returned.
    """

    def __init__(self, start: int = 0):
        self._pos = start
        self._parts: list[str] = []

    def consume(self, text: str) -> tuple[int, int]:
        start = self._pos
        end = start + len(text)
        self._parts.append(text)
        self._pos = end + 1  # +1 for the "\n" that join() will insert
        return start, end

    def skip(self) -> tuple[int, int]:
        """(pos, pos) for continuation cells that occupy no space in NER text."""
        return self._pos, self._pos

    def build_full_text(self) -> str:
        return "\n".join(self._parts)

    @property
    def position(self) -> int:
        return self._pos


@dataclass
class DocumentStructure:
    """Structured view of an extracted document.

    ``full_text`` is the NER input. Every element/cell with start/end offsets
    satisfies ``full_text[start:end] == text`` when the document was extracted
    via ``extract_document``.
    """

    full_text: str
    elements: list | None  # None for unstructured (PDF) files
    tables: list
    is_structured: bool  # True for DOCX, False for PDF/fallback

    def check_invariant(self) -> list[str]:
        """Return [] when all offsets are consistent with full_text; violations otherwise."""
        violations = []
        for elem in self.elements or []:
            if elem.get("type") == "table":
                continue
            start, end, text = (
                elem.get("start"),
                elem.get("end"),
                elem.get("text"),
            )
            if None in (start, end, text):
                continue
            actual = self.full_text[start:end]
            if actual != text:
                violations.append(
                    f"Element {elem.get('type')!r} at [{start}:{end}]: "
                    f"expected {text!r}, got {actual!r}"
                )
        for table in self.tables or []:
            for cell in table.get("cells", []):
                if cell.get("isMergedContinuation"):
                    continue
                start, end, text = (
                    cell.get("start"),
                    cell.get("end"),
                    cell.get("text", ""),
                )
                if None in (start, end):
                    continue
                actual = self.full_text[start:end]
                if actual != text:
                    violations.append(
                        f"Cell [{cell.get('row')},{cell.get('col')}] at [{start}:{end}]: "
                        f"expected {text!r}, got {actual!r}"
                    )
        return violations


def _paragraph_is_hidden(para) -> bool:
    """Return True if every text-bearing run has white (#FFFFFF) colour."""
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs:
        return False
    for run in text_runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return False
        color_el = rpr.find(qn("w:color"))
        if color_el is None:
            return False
        if color_el.get(qn("w:val"), "").upper() != "FFFFFF":
            return False
    return True


def _table_has_borders(table) -> bool:
    """Return True if any cell in the table has a visible border set."""
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return True

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return True

    border_names = [
        "w:top",
        "w:left",
        "w:bottom",
        "w:right",
        "w:insideH",
        "w:insideV",
    ]
    no_border_values = {"none", "nil"}

    for name in border_names:
        border_el = tbl_borders.find(qn(name))
        if border_el is not None:
            val = border_el.get(qn("w:val"))
            if val and val not in no_border_values:
                return True

    return False


def extract_table_with_styling(
    table, table_start_position: int, has_borders: bool = True
) -> tuple[str, list]:
    """Extract a DOCX table as structured data, preserving cell text and formatting."""
    html = '<table style="border-collapse: collapse; width: 100%;">'
    cells = []
    position = table_start_position

    # Pre-compute horizontal colspan
    row_tc_colspan = []
    for row in table.rows:
        tc_count = {}
        for cell in row.cells:
            tc = cell._tc
            tc_count[tc] = tc_count.get(tc, 0) + 1
        row_tc_colspan.append(tc_count)

    # Pre-compute vertical rowspan
    tc_first_row = {}
    tc_rowspan = {}
    for row_idx, row in enumerate(table.rows):
        seen_in_row = set()
        for cell in row.cells:
            tc = cell._tc
            if tc not in seen_in_row:
                seen_in_row.add(tc)
                if tc not in tc_first_row:
                    tc_first_row[tc] = row_idx
                    tc_rowspan[tc] = 1
                else:
                    tc_rowspan[tc] += 1

    for row_idx, row in enumerate(table.rows):
        html += "<tr>"
        seen_tc_in_row = set()

        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            is_h_continuation = tc in seen_tc_in_row
            is_v_continuation = tc_first_row.get(tc) != row_idx
            is_continuation = is_h_continuation or is_v_continuation
            seen_tc_in_row.add(tc)

            colspan = row_tc_colspan[row_idx].get(tc, 1)
            rowspan = tc_rowspan.get(tc, 1)
            cell_text = cell.text.replace("\n", " ")

            if is_continuation:
                cells.append(
                    {
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell_text,
                        "start": position,
                        "end": position,
                        "style": "",
                        "bgColor": None,
                        "runs": [],
                        "colspan": colspan,
                        "rowspan": rowspan,
                        "isMergedContinuation": True,
                    }
                )
                continue

            cell_start = position
            cell_end = position + len(cell_text)
            position = cell_end + 1

            tc_pr = cell._tc.get_or_add_tcPr()
            shading = tc_pr.find(qn("w:shd"))
            bg_color = None
            if shading is not None:
                fill = shading.get(qn("w:fill"))
                if fill and fill != "auto":
                    bg_color = f"#{fill}"

            cell_style = (
                "border: 1px solid #000; padding: 6px 8px;"
                if has_borders
                else "padding: 6px 8px;"
            )
            if bg_color:
                cell_style += f" background-color: {bg_color};"

            cell_html = ""
            text_runs = []
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if not text:
                        continue
                    styles = []
                    if run.bold:
                        styles.append("font-weight: bold")
                    if run.italic:
                        styles.append("font-style: italic")
                    if run.font.color and run.font.color.rgb:
                        styles.append(f"color: #{run.font.color.rgb}")

                    text_runs.append(
                        {
                            "text": text,
                            "style": "; ".join(styles) if styles else None,
                        }
                    )

                    if styles:
                        cell_html += (
                            f'<span style="{"; ".join(styles)}">{text}</span>'
                        )
                    else:
                        cell_html += text
                cell_html += "<br>"

            cells.append(
                {
                    "row": row_idx,
                    "col": col_idx,
                    "text": cell_text,
                    "start": cell_start,
                    "end": cell_end,
                    "style": cell_style,
                    "bgColor": bg_color,
                    "runs": text_runs,
                    "colspan": colspan,
                    "rowspan": rowspan,
                    "isMergedContinuation": False,
                }
            )

            span_attrs = ""
            if colspan > 1:
                span_attrs += f' colspan="{colspan}"'
            if rowspan > 1:
                span_attrs += f' rowspan="{rowspan}"'
            html += f'<td{span_attrs} style="{cell_style}">{cell_html.rstrip("<br>")}</td>'

        html += "</tr>"

    html += "</table>"
    return html, cells


def extract_document_structure(path: str):
    """Extract headings, paragraphs, and tables from a DOCX file.

    Returns (elements, tables, full_text) on success, or (None, None) for
    non-DOCX paths and parse failures.
    """
    if not path.lower().endswith(".docx"):
        return None, None

    try:
        doc = Document(path)
    except Exception as e:
        logger.error("Failed to parse DOCX at %s: %s", path, e)
        return None, None

    cursor = _PositionCursor()
    elements = []
    tables_data = []
    table_index = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text
                    if not text.strip():
                        continue
                    if _paragraph_is_hidden(para):
                        continue

                    style_name = para.style.name if para.style else "Normal"
                    element_type = (
                        "heading"
                        if style_name.startswith("Heading")
                        else "paragraph"
                    )
                    level = None
                    if element_type == "heading":
                        try:
                            level = int(style_name.split()[-1])
                        except (ValueError, IndexError):
                            level = 1

                    start, end = cursor.consume(text)
                    elements.append(
                        {
                            "type": element_type,
                            "level": level,
                            "text": text,
                            "start": start,
                            "end": end,
                        }
                    )
                    break

        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._tbl is element:
                    seen_tc_global = set()
                    text_rows = []
                    for row in tbl.rows:
                        row_cells = []
                        for cell in row.cells:
                            if cell._tc not in seen_tc_global:
                                seen_tc_global.add(cell._tc)
                                row_cells.append(cell.text.replace("\n", " "))
                        if row_cells:
                            text_rows.append("\t".join(row_cells))
                    table_text = "\n".join(text_rows)

                    has_borders = _table_has_borders(tbl)
                    table_start = cursor.position
                    html, cells = extract_table_with_styling(
                        tbl, table_start, has_borders
                    )
                    start, end = cursor.consume(table_text)

                    col_widths = []
                    try:
                        raw_widths = [col.width for col in tbl.columns]
                        total = sum(w for w in raw_widths if w)
                        if total:
                            col_widths = [
                                round(w / total * 100, 2) if w else None
                                for w in raw_widths
                            ]
                    except Exception as e:
                        logger.debug(
                            "Could not compute table column widths: %s", e
                        )

                    tables_data.append(
                        {
                            "id": table_index,
                            "html": html,
                            "text": table_text,
                            "cells": cells,
                            "hasBorders": has_borders,
                            "colWidths": col_widths,
                            "ner_start": start,
                            "ner_end": end,
                        }
                    )
                    elements.append(
                        {
                            "type": "table",
                            "table_id": table_index,
                            "start": start,
                            "end": end,
                        }
                    )
                    table_index += 1
                    break

    return elements, tables_data, cursor.build_full_text()


def _extract_text_from_pdf(path: str) -> str:
    """Extract plain text from a PDF file using pypdf."""
    try:
        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)
        return "\n\n".join(parts)
    except Exception as e:
        logger.error("Failed to extract text from PDF %s: %s", path, e)
        return ""


def extract_document(path: str) -> DocumentStructure:
    """Single entry point returning a DocumentStructure for any supported file type.

    Tries DOCX extraction first; falls back to PDF for all other paths or on
    parse failure.
    """
    result = extract_document_structure(path)
    if result[0] is not None:
        elements, tables, full_text = result
        return DocumentStructure(
            full_text=full_text,
            elements=elements,
            tables=tables,
            is_structured=True,
        )
    full_text = _extract_text_from_pdf(path)
    return DocumentStructure(
        full_text=full_text,
        elements=None,
        tables=[],
        is_structured=False,
    )
