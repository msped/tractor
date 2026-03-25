import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX

from cases.models import Case, Redaction
from cases.models import Document as CaseDocument
from ..models import Model, TrainingDocument, TrainingRun, TrainingRunCaseDoc, TrainingRunTrainingDoc
from ..tasks import (
    _build_spancat_pipeline,
    _prepare_examples,
    _run_training_loop,
    collect_training_data_detailed,
    train_model,
)
from .base import NetworkBlockerMixin

User = get_user_model()


def create_test_docx(path, highlights, multi_paragraph=False):
    """Helper to create a docx file with highlighted text.

    Args:
        path: Path to save the docx file
        highlights: List of (text, color) tuples or list of lists for multi-paragraph
        multi_paragraph: If True, highlights is a list of paragraphs, each containing
                        a list of (text, color) tuples
    """
    doc = DocxDocument()
    if multi_paragraph:
        for para_highlights in highlights:
            p = doc.add_paragraph()
            for text, color in para_highlights:
                run = p.add_run(text)
                if color:
                    run.font.highlight_color = color
    else:
        p = doc.add_paragraph()
        for text, color in highlights:
            run = p.add_run(text)
            if color:
                run.font.highlight_color = color
    doc.save(path)


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="test_media_collect"))
class CollectTrainingDataDetailedTests(NetworkBlockerMixin, TestCase):
    def setUp(self):
        self.user = User.objects.create_user("testuser", password="password")
        self.case = Case.objects.create(case_reference="C1")

        # Create a test .docx file for TrainingDocument
        self.docx_path = Path(tempfile.mkdtemp()) / "test.docx"
        create_test_docx(
            self.docx_path,
            [
                ("This is ", None),
                ("Third Party", WD_COLOR_INDEX.BRIGHT_GREEN),
                (" info. And this is ", None),
                ("Operational", WD_COLOR_INDEX.TURQUOISE),
                (" data.", None),
            ],
        )
        with open(self.docx_path, "rb") as f:
            self.training_doc = TrainingDocument.objects.create(
                name="Test Docx",
                original_file=SimpleUploadedFile(
                    "test.docx",
                    f.read(),
                    "application/vnd.openxmlformats-officedocument.\
                        wordprocessingml.document",
                ),
                created_by=self.user,
                processed=False,
            )

        # Create a CaseDocument with redactions
        self.case_doc = CaseDocument.objects.create(
            case=self.case,
            extracted_text="Some text with accepted redactions.",
            status=CaseDocument.Status.COMPLETED,
        )
        Redaction.objects.create(
            document=self.case_doc,
            start_char=15,
            end_char=23,
            text="accepted",
            is_accepted=True,
            redaction_type=Redaction.RedactionType.THIRD_PARTY_PII,
        )
        # This one should be ignored
        Redaction.objects.create(
            document=self.case_doc,
            start_char=0,
            end_char=4,
            text="Some",
            is_accepted=False,
            redaction_type=Redaction.RedactionType.THIRD_PARTY_PII,
        )

    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()

    def test_collect_from_training_docs(self):
        """Test collecting data only from TrainingDocuments."""
        train_data, t_docs, c_docs = collect_training_data_detailed(
            source="training_docs")

        self.assertEqual(len(train_data), 1)
        self.assertEqual(len(t_docs), 1)
        self.assertEqual(len(c_docs), 0)
        self.assertEqual(t_docs[0], self.training_doc)

        text, annotations = train_data[0]
        self.assertIn("Third Party", text)
        self.assertIn("Operational", text)
        self.assertEqual(len(annotations["entities"]), 2)
        self.assertEqual(annotations["entities"][0][2], "THIRD_PARTY")
        self.assertEqual(annotations["entities"][1][2], "OPERATIONAL")

    def test_collect_from_redactions(self):
        """Test collecting data only from CaseDocument redactions."""
        train_data, t_docs, c_docs = collect_training_data_detailed(
            source="redactions")

        self.assertEqual(len(train_data), 1)
        self.assertEqual(len(t_docs), 0)
        self.assertEqual(len(c_docs), 1)
        self.assertEqual(c_docs[0], self.case_doc)

        text, annotations = train_data[0]
        self.assertEqual(text, self.case_doc.extracted_text)
        self.assertEqual(len(annotations["entities"]), 1)
        start, end, label = annotations["entities"][0]
        self.assertEqual(start, 15)
        self.assertEqual(end, 23)
        self.assertEqual(label, "THIRD_PARTY")

    def test_collect_from_both(self):
        """Test collecting data from both sources."""
        train_data, t_docs, c_docs = collect_training_data_detailed(
            source="both")
        self.assertEqual(len(train_data), 2)
        self.assertEqual(len(t_docs), 1)
        self.assertEqual(len(c_docs), 1)


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="test_media_merge"))
class CollectTrainingDataMergeTests(NetworkBlockerMixin, TestCase):
    """Tests for merging adjacent highlighted runs."""

    def setUp(self):
        self.user = User.objects.create_user("testuser2", password="password")
        self.tmp_dir = tempfile.mkdtemp()

    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()

    def test_adjacent_runs_same_color_are_merged(self):
        """Test that adjacent runs with the same highlight color are merged into one entity."""
        docx_path = Path(self.tmp_dir) / "merge_test.docx"
        # Simulate Word splitting "Cheshire Police" into multiple runs
        create_test_docx(
            docx_path,
            [
                ("Cheshire ", WD_COLOR_INDEX.BRIGHT_GREEN),
                ("Police", WD_COLOR_INDEX.BRIGHT_GREEN),
                (" is here.", None),
            ],
        )
        with open(docx_path, "rb") as f:
            TrainingDocument.objects.create(
                name="Merge Test",
                original_file=SimpleUploadedFile("merge_test.docx", f.read()),
                created_by=self.user,
                processed=False,
            )

        train_data, t_docs, c_docs = collect_training_data_detailed(
            source="training_docs")

        self.assertEqual(len(train_data), 1)
        text, annotations = train_data[0]
        # Should be ONE entity for "Cheshire Police", not two separate ones
        self.assertEqual(len(annotations["entities"]), 1)
        start, end, label = annotations["entities"][0]
        self.assertEqual(text[start:end], "Cheshire Police")
        self.assertEqual(label, "THIRD_PARTY")

    def test_entity_boundaries_trimmed_of_whitespace(self):
        """Test that entity boundaries are trimmed to exclude leading/trailing whitespace."""
        docx_path = Path(self.tmp_dir) / "trim_test.docx"
        # Simulate highlighted text with surrounding whitespace
        create_test_docx(
            docx_path,
            [
                ("Text ", None),
                # whitespace on both sides
                (" Jones ", WD_COLOR_INDEX.BRIGHT_GREEN),
                (" more.", None),
            ],
        )
        with open(docx_path, "rb") as f:
            TrainingDocument.objects.create(
                name="Trim Test",
                original_file=SimpleUploadedFile("trim_test.docx", f.read()),
                created_by=self.user,
                processed=False,
            )

        train_data, t_docs, c_docs = collect_training_data_detailed(
            source="training_docs")

        self.assertEqual(len(train_data), 1)
        text, annotations = train_data[0]
        self.assertEqual(len(annotations["entities"]), 1)
        start, end, label = annotations["entities"][0]
        # Entity should be trimmed to just "Jones" without surrounding whitespace
        self.assertEqual(text[start:end], "Jones")

    def test_separate_highlights_remain_separate(self):
        """Test that highlights separated by non-highlighted text remain separate entities."""
        docx_path = Path(self.tmp_dir) / "separate_test.docx"
        create_test_docx(
            docx_path,
            [
                ("Name: ", None),
                ("John", WD_COLOR_INDEX.BRIGHT_GREEN),
                (" and ", None),
                ("Jane", WD_COLOR_INDEX.BRIGHT_GREEN),
                (" are here.", None),
            ],
        )
        with open(docx_path, "rb") as f:
            TrainingDocument.objects.create(
                name="Separate Test",
                original_file=SimpleUploadedFile(
                    "separate_test.docx", f.read()),
                created_by=self.user,
                processed=False,
            )

        train_data, t_docs, c_docs = collect_training_data_detailed(
            source="training_docs")

        self.assertEqual(len(train_data), 1)
        text, annotations = train_data[0]
        # Should be TWO separate entities
        self.assertEqual(len(annotations["entities"]), 2)
        entity_texts = [text[start:end]
                        for start, end, label in annotations["entities"]]
        self.assertIn("John", entity_texts)
        self.assertIn("Jane", entity_texts)

    def test_whitespace_only_highlights_are_ignored(self):
        """Test that highlights containing only whitespace are not included as entities."""
        docx_path = Path(self.tmp_dir) / "whitespace_test.docx"
        create_test_docx(
            docx_path,
            [
                ("Text ", None),
                # whitespace-only highlight
                ("   ", WD_COLOR_INDEX.BRIGHT_GREEN),
                (" more text ", None),
                ("Valid", WD_COLOR_INDEX.TURQUOISE),
            ],
        )
        with open(docx_path, "rb") as f:
            TrainingDocument.objects.create(
                name="Whitespace Test",
                original_file=SimpleUploadedFile(
                    "whitespace_test.docx", f.read()),
                created_by=self.user,
                processed=False,
            )

        train_data, t_docs, c_docs = collect_training_data_detailed(
            source="training_docs")

        self.assertEqual(len(train_data), 1)
        text, annotations = train_data[0]
        # Should be ONE entity (the whitespace-only one should be ignored)
        self.assertEqual(len(annotations["entities"]), 1)
        start, end, label = annotations["entities"][0]
        self.assertEqual(text[start:end], "Valid")


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="test_model_train"))
class TrainModelTests(NetworkBlockerMixin, TestCase):
    """Tests for the full SpanCat training pipeline."""

    def setUp(self):
        self.user = User.objects.create_user("trainuser", password="password")

    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()

    @patch("training.tasks.Task")
    def test_train_model_aborts_if_another_running(self, mock_task):
        """train_model() aborts if another task is already in progress."""
        mock_task.objects.filter.return_value.count.return_value = 1

        result = train_model(source="both")
        self.assertIsNone(result)

    @patch("training.tasks.Task")
    def test_train_model_aborts_with_insufficient_data(self, mock_task):
        """train_model() aborts if fewer than 25 training examples exist."""
        mock_task.objects.filter.return_value.count.return_value = 0

        initial_model_count = Model.objects.count()
        train_model(source="both")
        self.assertEqual(Model.objects.count(), initial_model_count)

    @patch("training.tasks._run_training_loop")
    @patch("training.tasks._prepare_examples")
    @patch("training.tasks._build_spancat_pipeline")
    @patch("training.tasks.collect_training_data_detailed")
    @patch("training.tasks.Task")
    def test_train_model_creates_spancat_model(
        self, mock_task, mock_collect, mock_build_pipeline, mock_prepare, mock_run_loop
    ):
        """train_model() creates a Model with model_type=SPANCAT and a TrainingRun."""
        mock_task.objects.filter.return_value.count.return_value = 0

        # Provide enough fake training data (>=25 entries)
        fake_data = [
            (f"text {i}", {"entities": [(0, 4, "THIRD_PARTY")]}) for i in range(25)]
        mock_collect.return_value = (fake_data, [], [])

        mock_nlp = MagicMock()
        mock_nlp.pipe_names = []
        mock_build_pipeline.return_value = mock_nlp
        mock_prepare.return_value = [MagicMock()]
        mock_run_loop.return_value = {
            "spans_sc_p": 0.8, "spans_sc_r": 0.75, "spans_sc_f": 0.77}

        initial_model_count = Model.objects.count()
        train_model(source="redactions")

        self.assertEqual(Model.objects.count(), initial_model_count + 1)
        new_model = Model.objects.order_by("-created_at").first()
        self.assertFalse(new_model.is_active)
        self.assertAlmostEqual(new_model.precision, 0.8)
        self.assertAlmostEqual(new_model.recall, 0.75)
        self.assertAlmostEqual(new_model.f1_score, 0.77)

        self.assertEqual(TrainingRun.objects.filter(
            model=new_model).count(), 1)

    @patch("training.tasks._run_training_loop")
    @patch("training.tasks._prepare_examples")
    @patch("training.tasks._build_spancat_pipeline")
    @patch("training.tasks.collect_training_data_detailed")
    @patch("training.tasks.Task")
    def test_train_model_cleans_up_on_exception(
        self, mock_task, mock_collect, mock_build_pipeline, mock_prepare, mock_run_loop
    ):
        """train_model() removes the output directory if training raises an exception."""
        mock_task.objects.filter.return_value.count.return_value = 0

        fake_data = [
            (f"text {i}", {"entities": [(0, 4, "THIRD_PARTY")]}) for i in range(25)]
        mock_collect.return_value = (fake_data, [], [])

        mock_build_pipeline.side_effect = RuntimeError("spaCy load failed")

        initial_model_count = Model.objects.count()
        with self.assertRaises(RuntimeError):
            train_model(source="redactions")

        # No model should be created
        self.assertEqual(Model.objects.count(), initial_model_count)

    @patch("training.tasks._run_training_loop")
    @patch("training.tasks._prepare_examples")
    @patch("training.tasks._build_spancat_pipeline")
    @patch("training.tasks.collect_training_data_detailed")
    @patch("training.tasks.Task")
    def test_train_model_records_training_and_case_docs(
        self, mock_task, mock_collect, mock_build_pipeline, mock_prepare, mock_run_loop
    ):
        """train_model() creates TrainingRunTrainingDoc and TrainingRunCaseDoc for used docs."""
        mock_task.objects.filter.return_value.count.return_value = 0

        case = Case.objects.create(case_reference="TRACK1")
        case_doc = CaseDocument.objects.create(
            case=case,
            extracted_text="Some case text",
            status=CaseDocument.Status.COMPLETED,
        )
        training_doc = TrainingDocument.objects.create(
            name="Track Doc",
            original_file=SimpleUploadedFile("track.docx", b"content"),
            created_by=self.user,
        )

        fake_data = [
            (f"text {i}", {"entities": [(0, 4, "THIRD_PARTY")]}) for i in range(25)]
        mock_collect.return_value = (fake_data, [training_doc], [case_doc])

        mock_nlp = MagicMock()
        mock_nlp.pipe_names = []
        mock_build_pipeline.return_value = mock_nlp
        mock_prepare.return_value = [MagicMock()]
        mock_run_loop.return_value = {
            "spans_sc_p": 0.8, "spans_sc_r": 0.75, "spans_sc_f": 0.77}

        train_model(source="both")

        new_model = Model.objects.order_by("-created_at").first()
        training_run = TrainingRun.objects.get(model=new_model)

        self.assertEqual(TrainingRunTrainingDoc.objects.filter(
            training_run=training_run).count(), 1)
        self.assertEqual(TrainingRunCaseDoc.objects.filter(
            training_run=training_run).count(), 1)


class BuildSpancatPipelineTests(NetworkBlockerMixin, TestCase):
    @patch("spacy.load")
    def test_builds_pipeline_with_spancat(self, mock_load):
        mock_nlp = MagicMock()
        mock_nlp.pipe_names = ["tok2vec", "ner", "parser"]
        mock_nlp.get_pipe.return_value.model.get_dim.return_value = 96
        mock_spancat = MagicMock()
        mock_nlp.add_pipe.return_value = mock_spancat
        mock_load.return_value = mock_nlp

        result = _build_spancat_pipeline()

        mock_load.assert_called_once_with("en_core_web_lg")
        mock_nlp.remove_pipe.assert_any_call("ner")
        mock_nlp.remove_pipe.assert_any_call("parser")
        mock_nlp.add_pipe.assert_called_once()
        self.assertEqual(mock_nlp.add_pipe.call_args[0][0], "spancat")
        mock_spancat.add_label.assert_any_call("THIRD_PARTY")
        mock_spancat.add_label.assert_any_call("OPERATIONAL")
        self.assertEqual(result, mock_nlp)


class PrepareExamplesTests(NetworkBlockerMixin, TestCase):
    def test_valid_spans_included(self):
        mock_nlp = MagicMock()
        mock_doc = MagicMock()
        mock_ref = MagicMock()
        mock_span = MagicMock()

        mock_nlp.make_doc.return_value = mock_doc
        mock_doc.copy.return_value = mock_ref
        mock_ref.char_span.return_value = mock_span

        train_data = [("Hello world", {"entities": [(0, 5, "THIRD_PARTY")]})]

        with patch("training.tasks.Example") as mock_example:
            mock_example.return_value = MagicMock()
            examples = _prepare_examples(mock_nlp, train_data)

        self.assertEqual(len(examples), 1)
        mock_ref.char_span.assert_called_once_with(
            0, 5, label="THIRD_PARTY", alignment_mode="contract")

    def test_misaligned_spans_dropped(self):
        mock_nlp = MagicMock()
        mock_doc = MagicMock()
        mock_ref = MagicMock()

        mock_nlp.make_doc.return_value = mock_doc
        mock_doc.copy.return_value = mock_ref
        mock_ref.char_span.return_value = None  # misaligned span

        train_data = [("Hello world", {"entities": [(0, 5, "THIRD_PARTY")]})]

        with patch("training.tasks.Example") as mock_example:
            mock_example.return_value = MagicMock()
            examples = _prepare_examples(mock_nlp, train_data)

        self.assertEqual(len(examples), 1)
        # The span was dropped, ref.spans["sc"] should be set to empty list
        mock_ref.spans.__setitem__.assert_called_with("sc", [])


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="test_media_train_loop"))
class RunTrainingLoopTests(NetworkBlockerMixin, TestCase):
    def test_early_stopping_triggered(self):
        mock_nlp = MagicMock()
        mock_nlp.resume_training.return_value = MagicMock()
        mock_nlp.evaluate.return_value = {
            "spans_sc_f": 0.0, "spans_sc_p": 0.0, "spans_sc_r": 0.0}

        with tempfile.TemporaryDirectory() as output_dir:
            _run_training_loop(mock_nlp, [], [], output_dir)

        # patience=3, so loop runs exactly 3 epochs before early stopping
        self.assertEqual(mock_nlp.evaluate.call_count, 3)

    def test_saves_to_disk_when_no_improvement(self):
        mock_nlp = MagicMock()
        mock_nlp.resume_training.return_value = MagicMock()
        mock_nlp.evaluate.return_value = {
            "spans_sc_f": 0.0, "spans_sc_p": 0.0, "spans_sc_r": 0.0}

        with tempfile.TemporaryDirectory() as output_dir:
            result = _run_training_loop(mock_nlp, [], [], output_dir)

        # When best_f1==0.0, nlp.to_disk is still called after the loop
        mock_nlp.to_disk.assert_called_once()
        self.assertEqual(result.get("spans_sc_f"), 0.0)


@override_settings(MEDIA_ROOT=tempfile.mkdtemp(prefix="test_media_edge"))
class CollectTrainingDataDetailedEdgeCasesTests(NetworkBlockerMixin, TestCase):
    def setUp(self):
        self.user = User.objects.create_user("edgeuser", password="password")

    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()

    def test_ds_information_redactions_excluded(self):
        """DS_INFORMATION redactions are excluded from training data."""
        from cases.models import Case, Redaction
        from cases.models import Document as CaseDocument

        case = Case.objects.create(case_reference="EDGE1")
        case_doc = CaseDocument.objects.create(
            case=case,
            extracted_text="DS info here and other stuff.",
            status=CaseDocument.Status.COMPLETED,
        )
        Redaction.objects.create(
            document=case_doc,
            start_char=0,
            end_char=7,
            text="DS info",
            is_accepted=True,
            redaction_type=Redaction.RedactionType.DS_INFORMATION,
        )

        train_data, _, _ = collect_training_data_detailed(source="redactions")
        # DS_INFORMATION excluded → entities empty → no training data for this doc
        self.assertEqual(len(train_data), 0)

    def test_completed_doc_with_no_text_skipped(self):
        """A completed document with empty extracted_text is skipped."""
        from cases.models import Case
        from cases.models import Document as CaseDocument

        case = Case.objects.create(case_reference="EDGE2")
        CaseDocument.objects.create(
            case=case,
            extracted_text="",
            status=CaseDocument.Status.COMPLETED,
        )

        train_data, _, _ = collect_training_data_detailed(source="redactions")
        self.assertEqual(len(train_data), 0)

    def test_bad_docx_logs_error_and_continues(self):
        """A corrupted training document is skipped without raising an exception."""
        from django.core.files.uploadedfile import SimpleUploadedFile

        bad_file = SimpleUploadedFile("bad.docx", b"not a valid docx file")
        TrainingDocument.objects.create(
            name="Bad Docx",
            original_file=bad_file,
            created_by=self.user,
            processed=False,
        )

        train_data, t_docs, _ = collect_training_data_detailed(
            source="training_docs")
        self.assertEqual(len(train_data), 0)
        self.assertEqual(len(t_docs), 0)
