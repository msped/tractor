import re

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

from .loader import GLiNERModelManager, SpanCatModelManager

_PREFIX_CHARS = {"#"}


def _paragraph_is_hidden(para):
    """Return True if every text-bearing run in the paragraph has white (#FFFFFF) colour.

    Word templates use white text to embed invisible content (e.g. form field markers,
    URN placeholders). These should be excluded from the extracted text.
    """
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs:
        return False
    for run in text_runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return False
        color_el = rpr.find(qn("w:color"))
        if color_el is None:
            return False
        if color_el.get(qn("w:val"), "").upper() != "FFFFFF":
            return False
    return True


def _expand_prefix_symbols(entities, text):
    """Extend entity start positions to include an immediately preceding '#' or similar symbol."""
    for ent in entities:
        start = ent["start_char"]
        if start > 0 and text[start - 1] in _PREFIX_CHARS:
            ent["start_char"] = start - 1
            ent["text"] = text[ent["start_char"] : ent["end_char"]]
    return entities


def _deduplicate_entities(primary_entities, secondary_entities):
    """Merge two entity lists, keeping primary entities and dropping overlapping secondary ones."""
    combined = list(primary_entities)
    for sec_ent in secondary_entities:
        overlaps = any(
            sec_ent["start_char"] < pri_ent["end_char"]
            and sec_ent["end_char"] > pri_ent["start_char"]
            for pri_ent in primary_entities
        )
        if not overlaps:
            combined.append(sec_ent)
    return combined


def _table_has_borders(table):
    """Return True if any cell in the table has a visible border set."""
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return True

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return True

    border_names = [
        "w:top",
        "w:left",
        "w:bottom",
        "w:right",
        "w:insideH",
        "w:insideV",
    ]
    no_border_values = {"none", "nil"}

    for name in border_names:
        border_el = tbl_borders.find(qn(name))
        if border_el is not None:
            val = border_el.get(qn("w:val"))
            if val and val not in no_border_values:
                return True

    return False


def extract_table_with_styling(table, table_start_position, has_borders=True):
    """Extract a DOCX table as structured data, preserving cell text and formatting."""
    html = '<table style="border-collapse: collapse; width: 100%;">'
    cells = []
    position = table_start_position

    # Pre-compute horizontal colspan: count how many times each _tc appears per row.
    # Use lxml elements directly as dict keys — their __hash__/__eq__ are based on the
    # underlying C-level node pointer, which is stable even if Python proxies are GC'd.
    row_tc_colspan = []
    for row in table.rows:
        tc_count = {}
        for cell in row.cells:
            tc = cell._tc
            tc_count[tc] = tc_count.get(tc, 0) + 1
        row_tc_colspan.append(tc_count)

    # Pre-compute vertical rowspan: count distinct rows each _tc appears in
    tc_first_row = {}
    tc_rowspan = {}
    for row_idx, row in enumerate(table.rows):
        seen_in_row = set()
        for cell in row.cells:
            tc = cell._tc
            if tc not in seen_in_row:
                seen_in_row.add(tc)
                if tc not in tc_first_row:
                    tc_first_row[tc] = row_idx
                    tc_rowspan[tc] = 1
                else:
                    tc_rowspan[tc] += 1

    for row_idx, row in enumerate(table.rows):
        html += "<tr>"
        seen_tc_in_row = set()

        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            is_h_continuation = tc in seen_tc_in_row
            is_v_continuation = tc_first_row.get(tc) != row_idx
            is_continuation = is_h_continuation or is_v_continuation
            seen_tc_in_row.add(tc)

            colspan = row_tc_colspan[row_idx].get(tc, 1)
            rowspan = tc_rowspan.get(tc, 1)

            cell_text = cell.text.replace("\n", " ")

            if is_continuation:
                # Continuation cells don't occupy space in the NER text
                # (deduplication mirrors extract_document_structure table_text)
                cells.append(
                    {
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell_text,
                        "start": position,
                        "end": position,
                        "style": "",
                        "bgColor": None,
                        "runs": [],
                        "colspan": colspan,
                        "rowspan": rowspan,
                        "isMergedContinuation": True,
                    }
                )
                continue

            # Advance position only for unique (non-continuation) cells
            cell_start = position
            cell_end = position + len(cell_text)
            position = cell_end + 1

            # Get cell shading (background color)
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = tc_pr.find(qn("w:shd"))
            bg_color = None
            if shading is not None:
                fill = shading.get(qn("w:fill"))
                if fill and fill != "auto":
                    bg_color = f"#{fill}"

            # Build cell style
            cell_style = (
                "border: 1px solid #000; padding: 6px 8px;"
                if has_borders
                else "padding: 6px 8px;"
            )
            if bg_color:
                cell_style += f" background-color: {bg_color};"

            # Get text with run formatting for HTML
            cell_html = ""
            text_runs = []
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if not text:
                        continue
                    styles = []
                    if run.bold:
                        styles.append("font-weight: bold")
                    if run.italic:
                        styles.append("font-style: italic")
                    if run.font.color and run.font.color.rgb:
                        styles.append(f"color: #{run.font.color.rgb}")

                    text_runs.append(
                        {
                            "text": text,
                            "style": "; ".join(styles) if styles else None,
                        }
                    )

                    if styles:
                        cell_html += (
                            f'<span style="{"; ".join(styles)}">{text}</span>'
                        )
                    else:
                        cell_html += text
                cell_html += "<br>"

            cells.append(
                {
                    "row": row_idx,
                    "col": col_idx,
                    "text": cell_text,
                    "start": cell_start,
                    "end": cell_end,
                    "style": cell_style,
                    "bgColor": bg_color,
                    "runs": text_runs,
                    "colspan": colspan,
                    "rowspan": rowspan,
                    "isMergedContinuation": False,
                }
            )

            span_attrs = ""
            if colspan > 1:
                span_attrs += f' colspan="{colspan}"'
            if rowspan > 1:
                span_attrs += f' rowspan="{rowspan}"'
            html += f'<td{span_attrs} style="{cell_style}">{cell_html.rstrip("<br>")}</td>'

        html += "</tr>"

    html += "</table>"
    return html, cells


def extract_document_structure(path):
    """Extract headings, paragraphs, and tables from a DOCX file as a structured list."""
    if not path.lower().endswith(".docx"):
        return None, None

    try:
        doc = Document(path)
    except Exception:
        return None, None

    elements = []
    tables_data = []
    position = 0
    full_text_parts = []
    table_index = 0

    # Iterate through document body elements in order
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # Remove namespace

        if tag == "p":
            # Find the matching paragraph object
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text
                    if not text.strip():
                        continue
                    if _paragraph_is_hidden(para):
                        continue

                    style_name = para.style.name if para.style else "Normal"
                    element_type = (
                        "heading"
                        if style_name.startswith("Heading")
                        else "paragraph"
                    )
                    level = None
                    if element_type == "heading":
                        try:
                            level = int(style_name.split()[-1])
                        except (ValueError, IndexError):
                            level = 1

                    elements.append(
                        {
                            "type": element_type,
                            "level": level,
                            "text": text,
                            "start": position,
                            "end": position + len(text),
                        }
                    )
                    full_text_parts.append(text)
                    position += len(text) + 1  # +1 for newline
                    break

        elif tag == "tbl":
            # Find the matching table object
            for tbl in doc.tables:
                if tbl._tbl is element:
                    # Build plain text for NER (tab-separated values).
                    # Use a global seen set so colspan/rowspan cells only
                    # contribute their text once across the entire table.
                    seen_tc_global = set()
                    text_rows = []
                    for row in tbl.rows:
                        row_cells = []
                        for cell in row.cells:
                            if cell._tc not in seen_tc_global:
                                seen_tc_global.add(cell._tc)
                                row_cells.append(
                                    cell.text.replace("\n", " ")
                                )
                        if row_cells:
                            text_rows.append("\t".join(row_cells))
                    table_text = "\n".join(text_rows)

                    # Check if table has visible borders
                    has_borders = _table_has_borders(tbl)

                    # Extract styled HTML and cell data with positions
                    html, cells = extract_table_with_styling(
                        tbl, position, has_borders
                    )

                    # Extract column widths from DOCX XML as percentages of total table width
                    col_widths = []
                    try:
                        raw_widths = [col.width for col in tbl.columns]
                        total = sum(w for w in raw_widths if w)
                        if total:
                            col_widths = [
                                round(w / total * 100, 2) if w else None
                                for w in raw_widths
                            ]
                    except Exception:
                        pass

                    tables_data.append(
                        {
                            "id": table_index,
                            "html": html,
                            "text": table_text,
                            "cells": cells,
                            "hasBorders": has_borders,
                            "colWidths": col_widths,
                            "ner_start": position,
                            "ner_end": position + len(table_text),
                        }
                    )

                    # Add table element to structure
                    elements.append(
                        {
                            "type": "table",
                            "table_id": table_index,
                            "start": position,
                            "end": position + len(table_text),
                        }
                    )

                    full_text_parts.append(table_text)
                    position += len(table_text) + 1
                    table_index += 1
                    break

    full_text = "\n".join(full_text_parts)
    return elements, tables_data, full_text


def _extract_text_from_pdf(path):
    """Extract plain text from a PDF file using pypdf."""
    try:
        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)
        return "\n\n".join(parts)
    except Exception:
        return ""


def extract_entities_from_text(path):
    """Run the full NLP pipeline (GLiNER → SpanCat → Presidio) on a document file.

    Returns:
        Tuple of (extracted_text, entity_suggestions, tables, spacy_model_entry).
    """
    # Lazy imports to avoid loading transformers/pandas at module import time
    from .extractors.gliner_extractor import extract_with_gliner
    from .extractors.presidio_extractor import (
        extract_operational_with_presidio,
        extract_with_presidio,
    )
    from .extractors.spancat_extractor import extract_with_spancat

    gliner_model = GLiNERModelManager.get_instance().get_model()
    if not gliner_model:
        raise ValueError("No GLiNER model available.")

    spancat_nlp = SpanCatModelManager.get_instance().get_model()

    # Try structure extraction for DOCX files first
    structure_result = extract_document_structure(path)

    if structure_result[0] is not None:
        # DOCX file - use python-docx extraction
        structure, tables, ner_text = structure_result

        if not ner_text or not ner_text.strip():
            return "", [], tables, structure

        gliner_results = extract_with_gliner(gliner_model, ner_text)
        presidio_tp = extract_with_presidio(ner_text)
        presidio_op = extract_operational_with_presidio(ner_text)
        spancat_results = (
            extract_with_spancat(spancat_nlp, ner_text) if spancat_nlp else []
        )

        # SpanCat > GLiNER > Presidio
        combined = _deduplicate_entities(spancat_results, gliner_results)
        combined = _deduplicate_entities(combined, presidio_tp + presidio_op)
        combined = _expand_prefix_symbols(combined, ner_text)

        return ner_text, combined, tables, structure

    # Fallback to pypdf for non-DOCX files (PDF, etc.)
    ner_text = _extract_text_from_pdf(path)

    # Normalize whitespace
    ner_text = re.sub(r"\n{3,}", "\n\n", ner_text)

    if not ner_text.strip():
        return "", [], [], None

    gliner_results = extract_with_gliner(gliner_model, ner_text)
    presidio_tp = extract_with_presidio(ner_text)
    presidio_op = extract_operational_with_presidio(ner_text)
    spancat_results = (
        extract_with_spancat(spancat_nlp, ner_text) if spancat_nlp else []
    )

    # SpanCat > GLiNER > Presidio
    combined = _deduplicate_entities(spancat_results, gliner_results)
    combined = _deduplicate_entities(combined, presidio_tp + presidio_op)
    combined = _expand_prefix_symbols(combined, ner_text)

    # For non-DOCX files, no tables and no structure (fallback to plain text rendering)
    return ner_text, combined, [], None
